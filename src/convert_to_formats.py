#!/usr/bin/env python3
"""Convierte INFORME_TECNICO_HERMES_WORD.md a .html y .docx"""

import markdown
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
import re
import os

base = '/home/usuario/word-hermes-prototype'
md_path = os.path.join(base, 'INFORME_TECNICO_HERMES_WORD.md')

with open(md_path, 'r', encoding='utf-8') as f:
    md_text = f.read()

# === 1. Convertir a HTML ===
html_body = markdown.markdown(md_text, extensions=['tables', 'fenced_code', 'toc'])

css = """
body { font-family: 'Segoe UI', -apple-system, sans-serif; max-width: 900px; margin: 40px auto; padding: 20px; color: #1a1a1a; line-height: 1.7; background: #fff; }
h1 { color: #e94560; border-bottom: 3px solid #e94560; padding-bottom: 10px; font-size: 28px; }
h2 { color: #0f3460; border-bottom: 2px solid #0f3460; padding-bottom: 6px; margin-top: 32px; font-size: 22px; }
h3 { color: #16213e; margin-top: 24px; font-size: 18px; }
h4 { color: #2a2a4a; font-size: 15px; }
code { background: #f4f4f4; padding: 2px 6px; border-radius: 3px; font-family: 'Cascadia Code', monospace; font-size: 0.9em; }
pre { background: #1a1a2e; color: #e0e0e0; padding: 16px; border-radius: 8px; overflow-x: auto; font-size: 13px; line-height: 1.5; }
pre code { background: none; padding: 0; color: inherit; }
table { border-collapse: collapse; width: 100%; margin: 16px 0; }
th { background: #0f3460; color: white; padding: 10px 14px; text-align: left; font-size: 13px; }
td { border: 1px solid #ddd; padding: 8px 14px; font-size: 13px; }
tr:nth-child(even) { background: #f9f9f9; }
blockquote { border-left: 4px solid #e94560; margin: 16px 0; padding: 8px 16px; background: #fef5f5; color: #555; }
hr { border: none; border-top: 2px solid #e0e0e0; margin: 32px 0; }
img { max-width: 100%; }
@media print { body { max-width: 100%; margin: 0; } }
"""

html_full = f"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Integracion Hermes Agent + Microsoft Word - Documento Tecnico</title>
<style>{css}</style>
</head>
<body>
{html_body}
</body>
</html>"""

html_path = os.path.join(base, 'INFORME_TECNICO_HERMES_WORD.html')
with open(html_path, 'w', encoding='utf-8') as f:
    f.write(html_full)
print(f'HTML: {html_path} ({len(html_full):,} chars)')

# === 2. Convertir a DOCX ===
doc = Document()

# Metadata
doc.core_properties.title = "Integracion Hermes Agent + Microsoft Word"
doc.core_properties.author = "Equipo Hermes Agent"
doc.core_properties.subject = "Documento Tecnico para Desarrollo"

# Margenes
section = doc.sections[0]
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# Header
header = section.header
hp = header.paragraphs[0]
hp.text = "Hermes Agent + Microsoft Word - Documento Tecnico"
hp.style.font.size = Pt(8)
hp.style.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# Footer
footer = section.footer
fp = footer.paragraphs[0]
fp.text = "Confidencial - Area de Desarrollo - Abril 2026"
fp.style.font.size = Pt(8)

lines = md_text.split('\n')
in_code_block = False
in_table = False

for line in lines:
    # Code blocks
    if line.startswith('```'):
        in_code_block = not in_code_block
        continue
    
    if in_code_block:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        continue
    
    # Table detection
    if line.startswith('|') and '---' not in line and not in_table:
        # Could be a table start - simplified: just add as paragraph for now
        cleaned = ' '.join([c.strip() for c in line.split('|') if c.strip()])
        if cleaned:
            p = doc.add_paragraph(cleaned)
            p.style.font.size = Pt(9)
        continue
    
    if line.startswith('|') and '---' in line:
        continue
    
    # Headings
    if line.startswith('# '):
        doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    elif line.startswith('#### '):
        doc.add_heading(line[5:], level=4)
    
    # Horizontal rule
    elif line.strip() == '---':
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(chr(0x2500) * 70)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    
    # Bullet list
    elif line.startswith('- ') or line.startswith('* '):
        doc.add_paragraph(line[2:], style='List Bullet')
    
    # Numbered list
    elif re.match(r'^\d+\.\s', line):
        doc.add_paragraph(re.sub(r'^\d+\.\s', '', line), style='List Number')
    
    # Empty
    elif line.strip() == '':
        pass
    
    # Regular paragraph
    else:
        p = doc.add_paragraph()
        # Handle inline code
        parts = re.split(r'(`[^`]+`)', line)
        for part in parts:
            if part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)
            else:
                # Handle bold
                bold_parts = re.split(r'(\*\*[^*]+\*\*)', part)
                for bp in bold_parts:
                    if bp.startswith('**') and bp.endswith('**'):
                        run = p.add_run(bp[2:-2])
                        run.bold = True
                    else:
                        p.add_run(bp)

docx_path = os.path.join(base, 'INFORME_TECNICO_HERMES_WORD.docx')
doc.save(docx_path)
print(f'DOCX: {docx_path}')

print()
print('3 formatos generados:')
for fmt in ['.md', '.html', '.docx']:
    fpath = base + '/INFORME_TECNICO_HERMES_WORD' + fmt
    size = os.path.getsize(fpath)
    print(f'  {fpath} ({size:,} bytes)')
